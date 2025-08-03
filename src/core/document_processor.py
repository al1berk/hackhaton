import os
from typing import Optional
from .ocr_processor import HandwritingOCR, is_image_file, supported_image_formats
from docx import Document


class DocumentProcessor:
    """Farklı formattaki dökümanları işleyen sınıf"""
    
    def __init__(self):
        self.ocr_processor = None  # Lazy loading
    
    def _init_ocr(self):
        """OCR processor'ı lazy loading ile başlat"""
        if self.ocr_processor is None:
            self.ocr_processor = HandwritingOCR()
        return self.ocr_processor
    
    @staticmethod
    def read_document(file_path: str) -> str:
        """Döküman dosyasını oku ve metin olarak döndür"""
        try:
            if not os.path.exists(file_path):
                print(f"❌ Dosya bulunamadı: {file_path}")
                return ""
            
            file_ext = os.path.splitext(file_path)[1].lower()
            print(f"📖 Dosya türü: {file_ext}")
            
            if file_ext == '.txt':
                return DocumentProcessor._read_txt(file_path)
            elif file_ext == '.pdf':
                return DocumentProcessor._read_pdf(file_path)
            elif file_ext == '.docx':
                return DocumentProcessor._read_docx(file_path)
            else:
                print(f"❌ Desteklenmeyen dosya formatı: {file_ext}")
                return ""
                
        except Exception as e:
            print(f"❌ Dosya okuma hatası: {e}")
            return ""
    
    @staticmethod
    def _read_txt(file_path: str) -> str:
        """TXT dosyasını oku"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            print(f"✅ TXT dosyası okundu: {len(content)} karakter")
            return content
        except UnicodeDecodeError:
            # Farklı encoding'leri dene
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    print(f"✅ TXT dosyası okundu ({encoding}): {len(content)} karakter")
                    return content
                except:
                    continue
            print("❌ TXT dosyası hiçbir encoding ile okunamadı")
            return ""
    
    @staticmethod
    def _read_pdf(file_path: str) -> str:
        """PDF dosyasını oku"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                
                print(f"📄 PDF sayfa sayısı: {len(reader.pages)}")
                
                for i, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        text += page_text + "\n"
                        print(f"✅ Sayfa {i+1} okundu")
                    except Exception as e:
                        print(f"⚠️  Sayfa {i+1} okunamadı: {e}")
                        continue
                
                print(f"✅ PDF dosyası okundu: {len(text)} karakter")
                return text
                
        except ImportError:
            print("❌ PyPDF2 kütüphanesi yüklü değil. Yüklemek için: pip install PyPDF2")
            return ""
        except Exception as e:
            print(f"❌ PDF okuma hatası: {e}")
            return ""
    
    @staticmethod
    def _read_docx(file_path: str) -> str:
        """DOCX dosyasını oku"""
        try:
            doc = Document(file_path)
            text = ""
            
            print(f"📄 DOCX paragraf sayısı: {len(doc.paragraphs)}")
            
            for i, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
            
            # Tabloları da oku
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            print(f"✅ DOCX dosyası okundu: {len(text)} karakter")
            return text
            
        except ImportError:
            print("❌ python-docx kütüphanesi yüklü değil. Yüklemek için: pip install python-docx")
            return ""
        except Exception as e:
            print(f"❌ DOCX okuma hatası: {e}")
            return ""
    
    @staticmethod
    def validate_file(file_path: str, max_size_mb: int = 10) -> bool:
        """Dosyayı validate et"""
        try:
            if not os.path.exists(file_path):
                print(f"❌ Dosya bulunamadı: {file_path}")
                return False
            
            # Dosya boyutu kontrolü
            file_size = os.path.getsize(file_path)
            max_size_bytes = max_size_mb * 1024 * 1024
            
            if file_size > max_size_bytes:
                print(f"❌ Dosya çok büyük: {file_size / (1024*1024):.1f}MB (Max: {max_size_mb}MB)")
                return False
            
            # Dosya uzantısı kontrolü
            supported_extensions = ['.txt', '.pdf', '.docx']
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext not in supported_extensions:
                print(f"❌ Desteklenmeyen dosya türü: {file_ext}")
                print(f"✅ Desteklenen türler: {', '.join(supported_extensions)}")
                return False
            
            print(f"✅ Dosya geçerli: {file_path} ({file_size / 1024:.1f}KB)")
            return True
            
        except Exception as e:
            print(f"❌ Dosya validasyon hatası: {e}")
            return False
    
    @staticmethod
    def get_file_info(file_path: str) -> dict:
        """Dosya hakkında bilgi al"""
        try:
            if not os.path.exists(file_path):
                return {"error": "Dosya bulunamadı"}
            
            stat = os.stat(file_path)
            file_ext = os.path.splitext(file_path)[1].lower()
            
            return {
                "name": os.path.basename(file_path),
                "path": file_path,
                "extension": file_ext,
                "size_bytes": stat.st_size,
                "size_mb": round(stat.st_size / (1024*1024), 2),
                "created": stat.st_ctime,
                "modified": stat.st_mtime
            }
        except Exception as e:
            return {"error": str(e)}
    
    def extract_text_from_image(self, file_path):
        """Görüntü dosyasından metin çıkar"""
        try:
            ocr = self._init_ocr()
            
            # Dosya boyutunu kontrol et
            file_size = os.path.getsize(file_path) / (1024 * 1024)  # MB
            if file_size > 10:  # 10MB limit
                return f"❌ Dosya çok büyük ({file_size:.1f}MB). Maksimum 10MB desteklenir."
            
            # OCR işlemi
            result = ocr.extract_mixed_text(file_path)
            
            if result['confidence'] == 'error':
                return f"❌ OCR hatası: {result.get('error', 'Bilinmeyen hata')}"
            
            text = result['text']
            if not text.strip():
                return "⚠️ Görüntüde metin bulunamadı."
            
            # Başarılı çıkarma
            info = f"\n\n📷 OCR Bilgileri:\n"
            info += f"🔍 Yöntem: {result['method']}\n"
            info += f"📊 Güven: {result['confidence']}\n"
            info += f"📝 Karakter sayısı: {len(text)}\n"
            
            return text + info
            
        except Exception as e:
            return f"❌ Görüntü işleme hatası: {str(e)}"
    
    def process_document(self, file_path):
        """Ana doküman işleme fonksiyonu - OCR desteği ile güncellenmiş"""
        if not os.path.exists(file_path):
            return "❌ Dosya bulunamadı."
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Mevcut format kontrolleri
        if file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        elif file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        
        # YENİ: Görüntü dosyası kontrolü
        elif is_image_file(file_path):
            return self.extract_text_from_image(file_path)
        
        else:
            supported_formats = ['.txt', '.pdf', '.docx', '.doc'] + supported_image_formats()
            return f"❌ Desteklenmeyen dosya formatı: {file_ext}\n" \
                   f"Desteklenen formatlar: {', '.join(supported_formats)}"